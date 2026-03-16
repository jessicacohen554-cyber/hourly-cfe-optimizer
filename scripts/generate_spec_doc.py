#!/usr/bin/env python3
"""Generate the Model Methodology & Specification Document as a .docx Word file."""

import os
import sys

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Constellation brand colors
BLUE = RGBColor(35, 114, 185)    # #2372B9
ORANGE = RGBColor(244, 123, 32)  # #F47B27
GREEN = RGBColor(107, 165, 67)   # #6BA543
GRAY = RGBColor(126, 128, 131)   # #7E8083
SEC_BLUE = RGBColor(0, 127, 164) # #007FA4
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)

def set_cell_shading(cell, hex_color):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): hex_color,
    })
    shading.append(shd)

def add_brand_bar(doc):
    """Add a thin colored horizontal rule using a 4-cell table."""
    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    colors = ['2372B9', 'F47B27', '6BA543', '007FA4']
    for i, color in enumerate(colors):
        cell = t.cell(0, i)
        cell.text = ''
        set_cell_shading(cell, color)
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.font.size = Pt(2)

def add_code_block(doc, code_text, font_size=8):
    """Add a code block with monospace font and gray background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(50, 50, 50)
    # Set paragraph shading
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F5F5F5',
    })
    pPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = WHITE
        set_cell_shading(cell, '2372B9')
    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def add_note_box(doc, text, color_hex='E8F4FD'):
    """Add an indented note/callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(80, 80, 80)
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color_hex,
    })
    pPr.append(shd)


def build_document():
    doc = Document()

    # ======================================================================
    # STYLES
    # ======================================================================
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(30, 30, 30)

    for level in range(1, 5):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Calibri'
        h_style.font.color.rgb = BLUE

    # ======================================================================
    # COVER PAGE
    # ======================================================================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('Hourly CFE Optimizer', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = BLUE
        run.font.size = Pt(28)

    subtitle = doc.add_heading('Model Methodology & Specification Document', level=0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.color.rgb = GRAY
        run.font.size = Pt(16)

    add_brand_bar(doc)

    meta_lines = [
        'Constellation Energy — Commercial Strategy & Analytics',
        'Document Version: 1.0 | Model Version: Pipeline v1.0.0',
        'Base Year: 2025 (Snapshot Model)',
        'Date: June 2025',
        'Classification: Internal — Confidential',
    ]
    for line in meta_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = GRAY

    doc.add_page_break()

    # ======================================================================
    # TABLE OF CONTENTS
    # ======================================================================
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Introduction',
        '3. Theoretical Framework',
        '4. Model Architecture',
        '   4.1 Data Engineering & Pipeline',
        '   4.2 Cost & Input Tables',
        '   4.3 Algorithm Selection & Core Mathematical Functions',
        '5. Validation Results',
        '6. Usage & Limitations',
        '7. Directions for Use',
        'Appendix A — Key Algorithm Code Blocks',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(10)

    doc.add_page_break()

    # ======================================================================
    # 1. EXECUTIVE SUMMARY
    # ======================================================================
    doc.add_heading('1. Executive Summary', level=1)

    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph(
        'The Hourly CFE Optimizer is a computational model that determines the least-cost portfolio '
        'of clean energy resources required to achieve any specified level of hourly clean energy '
        'matching (10%–99.99%) across seven major U.S. ISO/RTO regions. Unlike traditional annual '
        'procurement models, this optimizer evaluates resource mixes against 8,760 hours of actual '
        'demand and generation data, capturing the temporal mismatch between variable renewable '
        'generation and load that drives procurement costs at high clean energy targets.'
    )

    doc.add_heading('1.2 Scope', level=2)
    doc.add_paragraph(
        'The model spans seven ISOs representing approximately 70% of U.S. electricity consumption:'
    )
    add_table(doc,
        ['ISO', 'Demand (TWh)', 'Dimensions', 'Key Resources'],
        [
            ['CAISO', '224.0', '6D', 'Solar, wind, nuclear, offshore wind, geothermal'],
            ['ERCOT', '488.0', '4D', 'Solar, wind, nuclear'],
            ['PJM', '843.3', '5D', 'Nuclear, wind, solar, offshore wind'],
            ['NYISO', '151.6', '5D', 'Nuclear, hydro, wind, offshore wind'],
            ['NEISO', '115.3', '5D', 'Nuclear, hydro, wind, offshore wind'],
            ['MISO', '663.8', '4D', 'Nuclear, wind, solar'],
            ['SPP', '299.8', '4D', 'Wind, nuclear, solar'],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph('The pipeline produces three categories of output:')
    doc.add_paragraph(
        'Cost-optimized resource portfolios (Steps 1–2): For each (ISO, threshold, cost scenario), '
        'the minimum-cost resource mix that achieves the target hourly matching score, evaluated '
        'across up to 17,496 cost sensitivity combinations.',
        style='List Number'
    )
    doc.add_paragraph(
        'Marginal abatement cost (MAC) queues (Step 3): Path-dependent deployment sequences '
        'optimized for cheapest $/tCO₂ avoided, with resource lock-in and clean firm technology tranching.',
        style='List Number'
    )
    doc.add_paragraph(
        'Market simulation trajectories (Step 6 — SMARTargets): Forward-looking deployment simulations '
        'from 2023–2050 under reference, aspirational, and parametric emission reduction scenarios, '
        'incorporating Wright\'s Law learning curves, REC pricing, capacity markets, and LMP-driven revenue.',
        style='List Number'
    )

    doc.add_heading('1.3 Key Assumptions', level=2)
    assumptions = [
        ('2025 snapshot model', 'Generation profiles and grid mix reflect current conditions. Forward projections are modeled explicitly via demand growth rates and learning curves where applicable (Steps 2.2b, 3b, 6.1).'),
        ('ISO-level geographic resolution', 'Resources are sourced within each ISO/RTO region. No intra-ISO transmission constraints (copper-plate assumption). Transmission costs are flat $/MWh adders.'),
        ('Hydro is existing-only', 'No new hydroelectric capacity. Existing hydro available at wholesale market rates.'),
        ('No incrementality requirement', 'Buyers can claim existing clean generation via EAC procurement in baseline track.'),
        ('Perfect dispatch', 'No unit commitment constraints. Storage dispatch follows priority-ordered greedy algorithm (battery4 → battery8 → LDES → H₂).'),
        ('Load profile', 'Demand modeled using actual ISO-level 8,760-hour profiles from EIA-930.'),
    ]
    for title, desc in assumptions:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(title + ': ')
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(desc)
        run2.font.size = Pt(10)

    add_note_box(doc,
        '[Insert screenshot of docs/architecture-high-level.html here — High-Level System Architecture Diagram]'
    )

    doc.add_page_break()

    # ======================================================================
    # 2. INTRODUCTION
    # ======================================================================
    doc.add_heading('2. Introduction', level=1)

    doc.add_heading('2.1 Document Purpose', level=2)
    doc.add_paragraph(
        'This document provides a comprehensive specification of the Hourly CFE Optimizer model '
        'as implemented in the production codebase. It serves as:'
    )
    doc.add_paragraph('A technical manual for team members who will maintain, extend, or audit the model.', style='List Number')
    doc.add_paragraph('A methodology reference establishing the analytical basis, data provenance, algorithmic choices, and validation results.', style='List Number')
    doc.add_paragraph('A traceability record documenting how each pipeline step\'s outputs flow into downstream consumers.', style='List Number')

    doc.add_heading('2.2 Document Scope', level=2)
    doc.add_paragraph(
        'This specification covers pipeline Steps 0–6 (data ingestion through market simulation). '
        'Step 7 (dashboard data extraction) is excluded as it is a presentation layer subject to '
        'change and does not affect model outputs.'
    )

    doc.add_heading('2.3 How to Read This Document', level=2)
    doc.add_paragraph(
        'Section 3 establishes the theoretical framework and analytical basis. '
        'Section 4 provides the detailed model architecture, proceeding step-by-step through the '
        'pipeline with explicit references to functions, parameters, and data flows. Code blocks for '
        'key algorithms are collected in Appendix A and referenced inline. '
        'Section 5 covers validation, sensitivity analysis, and robustness checks. '
        'Section 6 documents limitations and edge cases. '
        'Section 7 provides practical usage instructions.'
    )

    doc.add_page_break()

    # ======================================================================
    # 3. THEORETICAL FRAMEWORK
    # ======================================================================
    doc.add_heading('3. Theoretical Framework', level=1)

    doc.add_heading('3.1 The 8,760-Hour Matching Problem', level=2)
    doc.add_paragraph(
        'The central analytical challenge is the temporal mismatch between variable renewable generation '
        'and electricity demand. Annual procurement accounting masks this mismatch — a buyer may claim '
        '100% clean energy annually while consuming fossil-generated power during nights, winters, and '
        'low-wind periods.'
    )
    doc.add_paragraph(
        'The emerging GHG Protocol Scope 2 revision (October 2025 first consultation draft) and SBTi '
        'Power Sector v2 framework (September 2025 draft) are moving toward hourly temporal matching. '
        'This model evaluates the cost and resource implications of this transition.'
    )
    doc.add_paragraph('Hourly matching score (HMS) is defined as:')
    add_code_block(doc, 'HMS = (1/8760) × Σ_h min(1, clean_supply(h) / demand(h))', font_size=10)
    doc.add_paragraph(
        'This score represents the fraction of demand met by clean energy in every hour, averaged across '
        'the year. A score of 95% means that in the average hour, 95% of demand is met by temporally '
        'coincident clean generation. This is strictly more demanding than annual matching.'
    )

    doc.add_heading('3.2 Resource Mix Optimization via Exhaustive Search', level=2)
    doc.add_paragraph(
        'Unlike linear programming (LP) or mixed-integer programming (MIP) approaches used by capacity '
        'expansion models such as GenX (Jenkins et al., 2017) or EPRI\'s US-REGEN, this model uses an '
        'exhaustive combinatorial search followed by cost-based selection. This approach was chosen because:'
    )
    doc.add_paragraph('The hourly matching constraint is non-convex (due to the min(1, ...) operator), making LP relaxations unreliable.', style='List Bullet')
    doc.add_paragraph('The search space (~1.6M combinations for 6D CAISO) is tractable with vectorized NumPy operations and memory-bounded chunking.', style='List Bullet')
    doc.add_paragraph('Exhaustive search guarantees global optimality within the grid resolution — no feasible mix is missed due to solver heuristics.', style='List Bullet')

    doc.add_heading('3.3 Marginal Abatement Cost (MAC) Framework', level=2)
    doc.add_paragraph(
        'The MAC queue (Step 3b) adopts a path-dependent sequential deployment framework where each '
        'threshold\'s resource deployment is constrained by ("ratcheted to") prior deployments. '
        'MAC is computed as:'
    )
    add_code_block(doc, 'MAC_t = Δ(NewBuildCost)_{t-1→t} / Δ(CO₂_Avoided)_{t-1→t}', font_size=10)
    doc.add_paragraph(
        'Costs include only new-build LCOE and transmission (no gas backup, no wholesale revenue). '
        'CO₂ avoided is computed from hourly dispatch-based fossil displacement with merit-order '
        'retirement (coal → oil → gas). This aligns with the World Bank\'s MAC curve methodology '
        'and is comparable to EPA\'s Integrated Planning Model (IPM) incremental cost approach.'
    )

    doc.add_heading('3.4 Market Simulation (SMARTargets)', level=2)
    doc.add_paragraph(
        'Step 6 employs a profit-driven deployment simulation where clean energy resources deploy '
        'wherever revenue exceeds cost. CFE level is an output (emerges from profitability), not an '
        'input constraint. Revenue is computed from endogenous LMP (hourly merit-order pricing), '
        'capacity markets, and scarcity-driven REC pricing. Cost includes Wright\'s Law deployment-based '
        'learning curves. This approach is closer to agent-based models of electricity markets than '
        'to traditional least-cost capacity expansion.'
    )

    doc.add_heading('3.5 Wright\'s Law Learning Curves', level=2)
    doc.add_paragraph(
        'Technology cost reductions follow Wright\'s Law (experience curves). Cost at cumulative '
        'deployment Q is:'
    )
    add_code_block(doc, 'C(Q) = C_FOAK × (Q / Q_ref)^(-b)\nwhere b = -log₂(1 - LR), LR = learning rate', font_size=10)
    doc.add_paragraph(
        'Technology-specific learning rates are sourced from published empirical literature. '
        'Solar ~20% (Swanson\'s Law), Battery Li-ion 18–20% (BloombergNEF 2024), '
        'Nuclear SMR 10–15% (DOE Liftoff 2023), CCS 10–12% (Global CCS Institute).'
    )

    doc.add_page_break()

    # ======================================================================
    # 4. MODEL ARCHITECTURE
    # ======================================================================
    doc.add_heading('4. Model Architecture', level=1)

    add_note_box(doc,
        '[Insert screenshot of docs/architecture-detailed.html here — Detailed Pipeline Architecture Diagram]'
    )

    # --- 4.0 Shared Modules ---
    doc.add_heading('4.0 Shared Analytical Modules', level=2)

    doc.add_heading('4.0.1 dispatch_utils.py — Hourly Dispatch Reconstruction', level=3)
    doc.add_paragraph(
        'This module provides the canonical dispatch algorithm used throughout the pipeline. '
        'All storage types carry state-of-charge (SOC) across window boundaries and apply round-trip '
        'efficiency per discharge event. The 4-phase storage dispatch order is:'
    )
    doc.add_paragraph('Battery 4-hour (Li-ion, 85% RTE, daily cycling)', style='List Number')
    doc.add_paragraph('Battery 8-hour (Li-ion, 85% RTE, daily cycling)', style='List Number')
    doc.add_paragraph('LDES 100-hour (iron-air, 50% RTE, 7-day rolling window)', style='List Number')
    doc.add_paragraph('H₂ 1000-hour (electrolysis + salt cavern + H₂ turbine, 35% RTE, 30-day window)', style='List Number')

    p = doc.add_paragraph()
    p.add_run('Key design decision: ').bold = True
    p.add_run(
        'Storage capacity in Steps 1–2.1 is expressed as a percentage of annual demand (energy capacity '
        'coefficient). In Step 3, capacity is translated back to physical units for dispatch: '
        'MW = coefficient × annual_MWh / duration_hours. See Appendix A, Code Block 1.'
    )

    doc.add_heading('4.0.2 pipeline_config.py — Single Source of Truth', level=3)
    doc.add_paragraph(
        'All constants, cost tables, resource caps, and parameters are defined in a single configuration '
        'module. Downstream scripts import from here — no local constant definitions are permitted.'
    )

    doc.add_heading('4.0.3 lmp_engine.py — Merit-Order LMP Pricing', level=3)
    doc.add_paragraph(
        'Constructs a fossil merit-order stack (coal → oil → gas with heat rate curves and emission costs) '
        'and computes hourly LMP as the marginal cost of the last fossil unit dispatched per hour.'
    )

    # --- 4.1 Data Engineering ---
    doc.add_heading('4.1 Data Engineering & Pipeline', level=2)

    doc.add_heading('4.1.1 Data Provenance & Cleaning', level=3)
    add_table(doc,
        ['Source', 'Files', 'Purpose'],
        [
            ['EIA-930', 'eia_generation_profiles.parquet, eia_demand_profiles.parquet', '8,760-hour demand & generation (2021–2025)'],
            ['EPA eGRID', 'egrid_emission_rates.json', 'Subregion emission factors by fuel'],
            ['NREL ATB 2024', 'Embedded in pipeline_config.py', 'LCOE tables, storage cost projections'],
            ['EIA AEO 2025', 'Embedded in pipeline_config.py', 'Demand growth rates (L/M/H)'],
            ['Lazard v17-18', 'Cross-validation', 'LCOE independent check'],
            ['LBNL Queued Up 2025', 'Embedded in TX tables', 'Transmission cost estimates'],
        ]
    )
    doc.add_paragraph()

    doc.add_paragraph(
        'Data transformation pipeline (eia_data_io.py): '
        '(1) Parquet ingestion from EIA-930 data. '
        '(2) Generation profile normalization: each fuel type\'s 8,760-hour profile normalized to sum=1.0. '
        '(3) Multi-year averaging (2021–2025) to reduce single-year weather bias — critical for wind (±15% interannual) and hydro (±25%). '
        '(4) Demand profile normalization: hourly demand normalized such that sum(normalized) = 8760. '
        '(5) Fossil mix profiles: hourly coal/gas/oil generation shares for merit-order dispatch.'
    )

    p = doc.add_paragraph()
    p.add_run('Time-series vs. static data usage:').bold = True
    doc.add_paragraph('Steps 1–2.1: Static 2025 demand and generation profiles.', style='List Bullet')
    doc.add_paragraph('Step 2.2 Phase 2: Time-varying demand by year via growth rates. Generation shapes remain 2025 (assumption: shape stable, magnitude scales with deployment).', style='List Bullet')
    doc.add_paragraph('Step 3b (MAC queue): Time-varying demand per SBTi year mapping (THRESHOLD_TARGET_YEARS).', style='List Bullet')
    doc.add_paragraph('Step 6 (SMARTargets): Time-varying demand per simulation year (2023–2050). LMP uses demand-scaled MW profiles.', style='List Bullet')

    # --- 4.2 Cost Tables ---
    doc.add_heading('4.2 Cost & Input Tables', level=2)

    doc.add_heading('4.2.1 Solar LCOE ($/MWh)', level=3)
    doc.add_paragraph('Source: NREL ATB 2024, regionalized using LBNL installed cost data.')
    add_table(doc,
        ['Level', 'CAISO', 'ERCOT', 'PJM', 'NYISO', 'NEISO', 'MISO', 'SPP'],
        [
            ['Low', '45', '40', '50', '70', '62', '48', '43'],
            ['Medium', '60', '54', '65', '92', '82', '62', '57'],
            ['High', '78', '70', '85', '120', '107', '82', '74'],
        ]
    )
    doc.add_paragraph(
        'Regional adjustment: Solar costs vary by irradiance (higher CF in ERCOT/SPP → lower LCOE), '
        'labor costs (higher in NYISO/NEISO), and permitting complexity. NYISO is 1.75× ERCOT due to '
        'NYC-area construction costs.'
    )

    doc.add_heading('4.2.2 Wind LCOE ($/MWh)', level=3)
    add_table(doc,
        ['Level', 'CAISO', 'ERCOT', 'PJM', 'NYISO', 'NEISO', 'MISO', 'SPP'],
        [
            ['Low', '55', '30', '47', '61', '55', '33', '28'],
            ['Medium', '73', '40', '62', '81', '73', '43', '37'],
            ['High', '95', '52', '81', '105', '95', '56', '48'],
        ]
    )
    doc.add_paragraph(
        'Regional adjustment: Wind LCOE driven by capacity factor (Class I–IV resources). '
        'SPP and ERCOT have Class I/II (40–50% CF). CAISO and NEISO have Class III/IV (25–35% CF).'
    )

    doc.add_heading('4.2.3 Offshore Wind LCOE ($/MWh)', level=3)
    add_table(doc,
        ['Level', 'CAISO', 'PJM', 'NYISO', 'NEISO'],
        [
            ['Low', '110', '65', '72', '68'],
            ['Medium', '150', '85', '95', '90'],
            ['High', '200', '112', '125', '118'],
        ]
    )
    doc.add_paragraph(
        'CAISO is dramatically higher due to floating technology. PJM cheapest fixed-bottom '
        '(shallowest water, NJ 7.5 GW mandate). ERCOT, MISO, SPP: no offshore resource.'
    )

    doc.add_heading('4.2.4 Nuclear New-Build LCOE ($/MWh)', level=3)
    add_table(doc,
        ['Level', 'CAISO', 'ERCOT', 'PJM', 'NYISO', 'NEISO', 'MISO', 'SPP'],
        [
            ['Low', '70', '68', '72', '75', '73', '70', '68'],
            ['Medium', '95', '90', '105', '110', '108', '100', '92'],
            ['High', '140', '135', '160', '170', '165', '155', '140'],
        ]
    )
    doc.add_paragraph('Source: NREL ATB 2024 SMR/advanced reactor estimates. Nuclear uprate LCOE: L=$15, M=$25, H=$40.')

    doc.add_heading('4.2.5 Storage Parameters', level=3)
    add_table(doc,
        ['Technology', 'Duration', 'RTE', 'Window', 'Source'],
        [
            ['Battery 4hr', '4 hours', '85%', 'Daily', 'NREL ATB 2024'],
            ['Battery 8hr', '8 hours', '85%', 'Daily', 'NREL ATB 2024'],
            ['LDES (iron-air)', '100 hours', '50%', '7-day rolling', 'DOE LDES Liftoff 2023'],
            ['Green H₂', '1,000 hours', '35%', '30-day rolling', 'Hydrogen Council 2024'],
        ]
    )
    doc.add_paragraph(
        'Storage costs expressed as annualized capacity cost per % of annual demand (not LCOS). '
        'Formula: price = CAPEX_per_kWh × (CRF + FOM_rate) × 1000 × regional_mult. '
        'CRF = 0.1019 (8% WACC, 20-year life).'
    )

    doc.add_heading('4.2.6 Wholesale Prices & Fuel Adjustments', level=3)
    add_table(doc,
        ['ISO', 'Wholesale ($/MWh)', 'Fuel Adj Low', 'Fuel Adj High'],
        [
            ['CAISO', '30', '-5', '+10'],
            ['ERCOT', '27', '-7', '+12'],
            ['PJM', '34', '-6', '+11'],
            ['NYISO', '42', '-4', '+8'],
            ['NEISO', '41', '-4', '+8'],
            ['MISO', '30', '-6', '+11'],
            ['SPP', '25', '-7', '+12'],
        ]
    )
    doc.add_paragraph('Source: EIA-930 + ISO annual market reports, 2024 weighted average DA LMP.')

    doc.add_heading('4.2.7 Demand Growth Rates (%/yr)', level=3)
    add_table(doc,
        ['ISO', 'Low', 'Medium', 'High'],
        [
            ['CAISO', '1.4%', '1.9%', '2.5%'],
            ['ERCOT', '2.0%', '3.5%', '5.5%'],
            ['PJM', '1.5%', '2.4%', '3.6%'],
            ['NYISO', '1.3%', '2.0%', '4.4%'],
            ['NEISO', '0.9%', '1.8%', '2.9%'],
            ['MISO', '1.2%', '2.2%', '3.8%'],
            ['SPP', '1.0%', '1.8%', '3.0%'],
        ]
    )
    doc.add_paragraph(
        'Source: EIA AEO 2025, NERC 2024 LTRA, ERCOT 2025 LTLF, PJM 2025 Load Forecast, '
        'Grid Strategies 2025. Low = baseline. Medium = confirmed large-load + moderate electrification. '
        'High = full data center/AI + accelerated electrification.'
    )

    # --- 4.3 Algorithm Details ---
    doc.add_heading('4.3 Algorithm Selection & Core Mathematical Functions', level=2)

    # STEP 1
    doc.add_heading('4.3.1 STEP 1.1a/b — Coarse Grid Sweep & Scoring', level=3)
    p = doc.add_paragraph()
    p.add_run('Scripts: ').bold = True
    p.add_run('step1_1a_generate_mixes.py, step1_1b_score_mixes.py')
    doc.add_paragraph(
        'Step 1.1a generates a Cartesian product of resource fractions at 5-percentage-point step '
        'for each ISO\'s resource dimensions using generate_resource_combos(iso, step=5) in '
        'step1_pfs_generator.py. For a 4D ISO (ERCOT): ~12,000 combinations. For 6D CAISO: ~1.6 million. '
        'Seed combos from prior research are added via get_seed_combos(iso). '
        'When prior windows are available, the Cartesian product is narrowed to union bounds ± 15pp '
        'plus 100 scout mixes, saving ~30% of the search space.'
    )
    doc.add_paragraph(
        'Step 1.1b loads EIA demand + generation profiles and scores all mixes via batch_hourly_scores() '
        'using vectorized 8,760-hour matching. Scoring is memory-bounded in chunks of 20,000 mixes '
        '(peak ~1.4 GiB). Each mix\'s clean supply profile is constructed by weighting each resource\'s '
        'normalized generation shape by its allocation percentage.'
    )
    p = doc.add_paragraph()
    p.add_run('Output: ').bold = True
    p.add_run('{ISO}_coarse_cache.parquet')

    doc.add_heading('4.3.2 STEP 1.2 — Zone-Based Fine Search', level=3)
    p = doc.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run('step1_2_zone_search.py')
    doc.add_paragraph(
        'Divides score space into three zones (A: 50–70%, B: 70–90%, C: 90–100%). For each zone: '
        '(1) Identify coarse boundary mixes. (2) Compute resource windows. (3) Generate 1% fine grid '
        '(FINE_STEP=1). (4) Deduplicate against global hash set (collision-free int64: key = Σ(round(resource_i) × 301^i)). '
        '(5) Score via batch_hourly_scores(). (6) Assign to thresholds + dominance filter. '
        'Safety caps: MAX_FINE_ARCHETYPES=2,000 (4D), 500 (5D+). Fallback to archetype expansion if >10M combos.'
    )
    p = doc.add_paragraph()
    p.add_run('Output: ').bold = True
    p.add_run('{ISO}_t{T}_raw_pfs.parquet (per threshold), {ISO}_near_miss.parquet (for Step 1.5)')

    doc.add_heading('4.3.3 STEP 1.3 — Floor-Aware PFS', level=3)
    p = doc.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run('step1_3_floor_aware_pfs.py')
    doc.add_paragraph(
        'Generates incremental resource additions above existing clean floor (GRID_MIX_SHARES). '
        'Grid: solar 0–80% (2% step), wind 0–80% (2% step), clean firm 0–40% (2% step), '
        'hydro fixed at existing, offshore wind 0–30% (5% step), geothermal 0–20% (5% step, CAISO). '
        'Targets 50–70%. Produces minimal new-build mixes critical for MAC accuracy.'
    )

    doc.add_heading('4.3.4 STEP 1.4 — Fine Grid PFS', level=3)
    p = doc.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run('step1_4_fine_grid_pfs.py')
    doc.add_paragraph(
        'Fills coverage gap for 40–70% thresholds using 1% grid (vs. Step 1.3\'s 2%). '
        'Solar 0–60%, wind 0–60%, clean firm 0–30%, offshore 0–20% (2% step), geo 0–15% (3% step).'
    )

    doc.add_heading('4.3.5 STEP 1.5 — Storage Refinement ★', level=3)
    p = doc.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run('step1_5_storage_refinement.py')
    doc.add_paragraph(
        'This step identifies resource mixes that fail via generation alone but can reach targets '
        'with storage dispatch. Uses a three-pass adaptive funnel:'
    )
    p = doc.add_paragraph()
    p.add_run('Pass 0 — Maximum Screen (~40s/ISO): ').bold = True
    p.add_run(
        'Score each near-miss mix with ceiling storage (bat4=0.10%, bat8=0.15%, LDES=1.0%, H₂=1.0% '
        'of annual demand). Eliminates mixes that cannot reach any threshold even with max storage.'
    )
    p = doc.add_paragraph()
    p.add_run('Pass 1 — Adaptive Coarse Sweep: ').bold = True
    p.add_run(
        'Group surviving mixes by gap-to-threshold into buckets (0–5pp, 5–10pp, 10–25pp, 25–50pp). '
        'Each bucket gets right-sized storage grid. Cartesian product across 4 storage dimensions. '
        'See Appendix A, Code Block 2 for grid definition.'
    )
    p = doc.add_paragraph()
    p.add_run('Pass 2 — Fine Targeted (0.05% resolution): ').bold = True
    p.add_run('Refine mixes near each threshold\'s storage-enhanced boundary.')
    p = doc.add_paragraph()
    p.add_run('Dominance filter: ').bold = True
    p.add_run(
        'After each pass, mixes where all resource and storage allocations exceed those of a passing '
        'mix are eliminated. Floor/fine augmentation loads Step 1.3/1.4 outputs and runs storage sweep.'
    )
    p = doc.add_paragraph()
    p.add_run('Output: ').bold = True
    p.add_run('{ISO}_t{T}_storage.parquet (per threshold)')

    # STEP 2
    doc.add_heading('4.3.6 STEP 2.1 — Efficient Frontier Extraction', level=3)
    p = doc.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run('step2_1_efficient_frontier.py')
    doc.add_paragraph(
        'Three-phase reduction: (1) Threshold gate — keep rows whose scores fall in target ranges. '
        '(2) Resource cap filter — enforce solar cap (100%), total procurement cap (350%), hydro cap '
        '(existing levels). (3) Global deduplication — for each unique allocation, keep highest score. '
        'No dominance removal across mixes — different mixes at same score can have different costs. '
        'Output partitioned into non-overlapping threshold bands by score.'
    )
    p = doc.add_paragraph()
    p.add_run('Traceability: ').bold = True
    p.add_run('Step 2.1 outputs → Step 2.2a (cost optimization) and Step 2.2b (track evaluation).')

    doc.add_heading('4.3.7 STEP 2.2a — Cost Optimization', level=3)
    p = doc.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run('step2_2a_cost_optimization.py')
    doc.add_paragraph(
        'For each (ISO, threshold): loads EF mixes with score ≥ threshold. Evaluates across all '
        'sensitivity combinations (5,832 non-CAISO; 17,496 CAISO). 9-dimension sensitivity key: '
        '{Ren}{Firm}{Batt}{LDES}_{Fuel}_{Tx}_{CCS}{45Q}_{Geo}. Vectorized cost evaluation using '
        'NumPy broadcasting. Selects cheapest mix per sensitivity combo.'
    )
    p = doc.add_paragraph()
    p.add_run('Phase 2 — Demand Growth Sweep: ').bold = True
    p.add_run(
        'Extract unique winning archetypes. For each (year, growth level): evaluate with demand scaled '
        'by compound growth and costs adjusted via Wright\'s Law learning curves.'
    )
    p = doc.add_paragraph()
    p.add_run('Traceability: ').bold = True
    p.add_run(
        'Step 2.2a → Step 3a (dispatch cache), Step 3b (MAC queue), Step 4 (analytics), Step 6 (SMARTargets).'
    )

    doc.add_heading('4.3.8 STEP 2.2b — Track NB/CTR', level=3)
    p = doc.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run('step2_2b_track_nb_ctr.py')
    doc.add_paragraph(
        'New-Build Track: only new resources (no existing credit). Cost-to-Replace Track: replace '
        'existing nuclear with new clean. Demand growth sweep with Wright\'s Law learning.'
    )
    p = doc.add_paragraph()
    p.add_run('Traceability: ').bold = True
    p.add_run('Step 2.2b → Step 4.1e (export tracks) → Step 4.2c (analyze tracks) → dashboard.')

    # STEP 3
    doc.add_heading('4.3.9 STEP 3a — Dispatch Cache', level=3)
    p = doc.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run('step3a_build_dispatch_cache.py')
    doc.add_paragraph(
        'Pre-computes full 8,760-hour dispatch for every unique mix from Step 2.2 using '
        'dispatch_utils.reconstruct_hourly_dispatch(detailed=True). Battery capacity translation '
        'from % of annual demand to physical MW: MW = (pct/100) × annual_MWh / duration_hours. '
        'See Appendix A, Code Block 1.'
    )
    p = doc.add_paragraph()
    p.add_run('Traceability: ').bold = True
    p.add_run('Step 3a → Step 4.1a (CO₂/LMP), Step 4.1b (day profiles), Step 4.2b (storage analysis), Step 5.2a (scenario comparison).')

    doc.add_heading('4.3.10 STEP 3b — MAC Queue', level=3)
    p = doc.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run('step3b_mac_queue.py')
    doc.add_paragraph(
        'Path-dependent deployment queue: 15 pathways per ISO (3 demand × 5 price sensitivity). '
        'At each threshold: (1) compute demand at SBTi year, (2) dispatch floor → baseline CO₂, '
        '(3) sample floor-respecting archetypes, (4) score as new_build_cost/CO₂_avoided, '
        '(5) winner = argmin(MAC) with ≤1% overshoot, (6) ratchet floor.'
    )
    p = doc.add_paragraph()
    p.add_run('Clean firm tranching (merit-order): ').bold = True
    p.add_run(
        'Tranche 1: Nuclear uprates (cheapest, capped). Tranche 2: Geothermal (CAISO, capped at 39 TWh). '
        'Tranche 3: min(nuclear new-build, CCS) — CCS capped per ISO. See Appendix A, Code Block 3.'
    )
    p = doc.add_paragraph()
    p.add_run('Traceability: ').bold = True
    p.add_run('Step 3b → Step 5.2a (scenario_a files + queue), Step 5.2b (strategy 1 MAC ordering).')

    # STEP 4
    doc.add_heading('4.3.11 STEP 4 — Derived Analytics', level=3)
    doc.add_paragraph(
        'Step 4.1a (fossil dispatch): Merit-order fossil stack from Step 3a cache. CO₂ and LMP computation. '
        'Step 4.1b (compressed day profiles): 8,760h → 24h representative day. '
        'Step 4.1c (MAC stats): PCHIP spline MAC curves, 9 crossover points, no-regrets analysis. '
        'Step 4.1d (optimal targets): Where grid MAC = DAC cost. '
        'Step 4.1e→4.2c: Export tracks → analyze P10/P50/P90 cost envelopes. '
        'Step 4.2a: Resource density strips for abatement dashboard. '
        'Step 4.2b: Battery/LDES dispatch metrics, seasonal patterns.'
    )

    # STEP 5
    doc.add_heading('4.3.12 STEP 5 — Procurement Strategy Evaluation', level=3)
    doc.add_paragraph(
        'Step 5.1 (Scenario B — Hourly Matching): Four supply pools: '
        '(1) SSS at $0 (policy-supported nuclear, hydro, RPS). (2) Contracted (excluded). '
        '(3) Existing merchant at EAC premium ($3–5/MWh). (4) New-build from hourly gap. '
        'Learning: FOAK(H)→NOAK(L) by 2040.'
    )
    doc.add_paragraph(
        'Step 5.2a (Scenario Comparison): Reads Scenario A from data/step3-dispatch/mac_queue/ '
        'and Scenario B from data/step5-scenarios/. Side-by-side cost, mix, gas, MAC trajectories.'
    )
    doc.add_paragraph(
        'Step 5.2b (Strategy 1 — Consequential): Three variants (grid-average, fossil-average, marginal baseline). '
        'Cross-regional. Step 5.2c (Strategy 2 — Hourly): Three variants (100% new-build, grid baseline, SSS+premium). '
        'Same-ISO. Step 5.2d (Strategy 3 — Annual bundled). Step 5.2e (Wright\'s Law curves export).'
    )

    # STEP 6
    doc.add_heading('4.3.13 STEP 6 — SMARTargets Market Simulation', level=3)
    p = doc.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run('step6_1_smartargets.py')
    doc.add_paragraph(
        'Market-driven clean energy deployment simulation across 7 ISOs, 2023–2050.'
    )
    p = doc.add_paragraph()
    p.add_run('Scenarios: ').bold = True
    p.add_run(
        'R1/R2 (Reference): Pure profit-driven. '
        'AT1–AT4 (Aspirational Transition): Profit + mandated emission constraint. '
        'QT1–QT4 (Quick Transition): Parametric sweep (19 reduction targets).'
    )
    p = doc.add_paragraph()
    p.add_run('Revenue model: ').bold = True
    p.add_run(
        '(1) Energy: generation-weighted LMP from lmp_engine.py. '
        '(2) Capacity market: ISO-specific, ELCC-degraded. '
        '(3) REC: scarcity-driven compliance pricing — price = ACP × (1 − exp(−k × gap%)). '
        'See Appendix A, Code Block 4.'
    )
    p = doc.add_paragraph()
    p.add_run('Cost model: ').bold = True
    p.add_run(
        'LCOE from pipeline_config + Wright\'s Law deployment-based learning (Appendix A, Code Block 5) '
        '+ transmission + PPA discount (regional market depth scaling).'
    )
    p = doc.add_paragraph()
    p.add_run('Parametric sweep: ').bold = True
    p.add_run('270 scenarios = 2 conditions × 3 demand × 5 price × 3 PPA × 3 gas friction.')

    doc.add_page_break()

    # ======================================================================
    # 5. VALIDATION
    # ======================================================================
    doc.add_heading('5. Validation Results', level=1)

    doc.add_heading('5.1 Physics Validation', level=2)
    doc.add_paragraph('Hourly matching scores: Existing grid mixes produce scores consistent with eGRID clean energy shares (within ±2pp).', style='List Bullet')
    doc.add_paragraph('Storage dispatch conservation: Total energy out ≤ total energy in × RTE for all storage types. SOC non-negative at all hours.', style='List Bullet')
    doc.add_paragraph('Hydro cap enforcement: No feasible mix exceeds ISO-specific caps after Step 2.1 filtering.', style='List Bullet')

    doc.add_heading('5.2 Cost Validation', level=2)
    doc.add_paragraph('Battery cross-check: 0.01% bat4 at CAISO, Medium = $4.16/MWh (model) vs. $4.13/MWh (physical calc). Error: 0.7%.', style='List Bullet')
    doc.add_paragraph('LCOE benchmarks: Model tables validated against NREL ATB 2024 medium case within ±5%.', style='List Bullet')

    doc.add_heading('5.3 Sensitivity Analysis', level=2)
    doc.add_paragraph(
        'The 9-dimensional sensitivity sweep (5,832–17,496 combinations per threshold) systematically '
        'varies renewable LCOE, firm clean LCOE, battery cost, LDES cost, fuel prices, transmission, '
        'CCS availability, 45Q credit, and geothermal. P10/P50/P90 cost envelopes capture full uncertainty range.'
    )

    doc.add_heading('5.4 MAC Validation', level=2)
    doc.add_paragraph('$0–50/tCO₂ for 50–70% clean (wind/solar) — consistent with BNEF LCOE tracker.', style='List Bullet')
    doc.add_paragraph('$100–300/tCO₂ for 90–95% clean (firm deployment) — consistent with Princeton Net Zero America.', style='List Bullet')
    doc.add_paragraph('$500+/tCO₂ for 99%+ (last-mile) — consistent with DOE Liftoff estimates.', style='List Bullet')

    doc.add_page_break()

    # ======================================================================
    # 6. USAGE & LIMITATIONS
    # ======================================================================
    doc.add_heading('6. Usage & Limitations', level=1)

    doc.add_heading('6.1 How to Interpret Output', level=2)
    doc.add_paragraph(
        'Cost results represent the theoretical minimum-cost resource portfolio under specified assumptions. '
        'Real-world procurement involves additional factors — developer availability, interconnection timelines, '
        'contract structure, counterparty risk.'
    )

    doc.add_heading('6.2 Known Limitations', level=2)
    limitations = [
        ('Static supply model', 'Does not account for price-induced supply responses.'),
        ('No cross-ISO interactions', 'Each ISO modeled independently.'),
        ('No intra-ISO transmission constraints', 'Copper-plate assumption. GenX (DCOPF) and US-REGEN (zonal pipe-and-bubble) capture congestion-driven price separation.'),
        ('No unit commitment constraints', 'Perfect dispatch assumed. No min up/down, ramp rates, start-up costs.'),
        ('No hourly reserves', 'Resource adequacy enforced, but not spinning/non-spinning/regulation reserves.'),
        ('No demand-side flexibility', 'Load perfectly inelastic. No DR, load shifting, or flexible consumption.'),
        ('Single-sector scope', 'Electricity-only. No cross-sector interactions.'),
        ('Policy evolution', 'Reflects early 2025 policy landscape.'),
        ('No interconnection queue constraints', 'New capacity assumed buildable (except SMARTargets queue caps).'),
    ]
    for title, desc in limitations:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(title + ': ')
        r.bold = True
        r.font.size = Pt(10)
        r2 = p.add_run(desc)
        r2.font.size = Pt(10)

    doc.add_heading('6.3 Future Enhancements', level=2)
    doc.add_paragraph(
        'Tier 1 (High Impact): Zonal transmission modeling, operating reserve & unit commitment, '
        'demand-side flexibility. Informed by EPRI US-REGEN and GenX.jl.'
    )
    doc.add_paragraph(
        'Tier 2 (Medium Impact): Cross-ISO EAC trade, endogenous capacity retirement, flexible CCS dispatch.'
    )
    doc.add_paragraph(
        'Tier 3 (Aspirational): DC optimal power flow, multi-stage pathway optimization, multi-sector integration.'
    )

    doc.add_page_break()

    # ======================================================================
    # 7. DIRECTIONS
    # ======================================================================
    doc.add_heading('7. Directions for Use', level=1)

    doc.add_heading('7.1 Environment Setup', level=2)
    add_code_block(doc, 'pip install numpy numba pyarrow pandas\npython3 -c "from numba import njit; print(\'Numba OK\')"')

    doc.add_heading('7.2 Running the Pipeline', level=2)
    add_code_block(doc, '''# Step 1: Physics Feasible Space (~3-8 hours)
python scripts/step1_1a_generate_mixes.py --iso ALL
python scripts/step1_1b_score_mixes.py --iso ALL
python scripts/step1_2_zone_search.py --iso ALL
python scripts/step1_3_floor_aware_pfs.py --iso ALL
python scripts/step1_4_fine_grid_pfs.py --iso ALL
python scripts/step1_5_storage_refinement.py --iso ALL

# Step 2: Optimization (~15-30 min)
python scripts/step2_1_efficient_frontier.py --iso ALL
python scripts/step2_2a_cost_optimization.py --iso ALL

# Step 3: Dispatch & MAC (~10-30 min)
python scripts/step3a_build_dispatch_cache.py
python scripts/step3b_mac_queue.py

# Step 4: Analytics (parallel, ~5-10 min)
python scripts/step4_1a_fossil_dispatch.py
python scripts/step4_1c_compute_mac_stats.py

# Step 5: Scenarios (~10-20 min)
python scripts/step5_1_scenario_hourly.py
python scripts/step5_2a_scenario_comparison.py

# Step 6: SMARTargets (20-60 min)
python scripts/step6_1_smartargets.py''')

    doc.add_heading('7.3 Output Locations', level=2)
    add_table(doc,
        ['Step', 'Directory', 'Format'],
        [
            ['0', 'data/eia-930/', 'Parquet'],
            ['1', 'data/step1-pfs/', 'Parquet'],
            ['2.1', 'data/step2.1-ef/', 'Parquet'],
            ['2.2', 'data/step2.2-cost/', 'Parquet'],
            ['3', 'data/step3-dispatch/', 'Parquet + JSON'],
            ['4', 'data/step4-analysis/', 'Parquet + JSON'],
            ['5', 'data/step5-scenarios/', 'JSON'],
            ['6', 'data/step6-smartargets/', 'Parquet + JSON'],
        ]
    )

    doc.add_page_break()

    # ======================================================================
    # APPENDIX A
    # ======================================================================
    doc.add_heading('Appendix A — Key Algorithm Code Blocks', level=1)

    doc.add_heading('Code Block 1: Hourly Dispatch Reconstruction (dispatch_utils.py)', level=2)
    doc.add_paragraph(
        'Core 4-phase storage dispatch function. All storage types carry SOC across window boundaries '
        'with round-trip efficiency per discharge event. This is the single source of truth for dispatch '
        'logic used by Steps 1, 3, 4, and 6.'
    )
    add_code_block(doc, '''def reconstruct_hourly_dispatch(demand_norm, supply_profiles, resource_pcts,
                                 procurement_pct, battery_dispatch_pct=0,
                                 battery8_dispatch_pct=0, ldes_dispatch_pct=0,
                                 h2_dispatch_pct=0, detailed=False):
    """Reconstruct 8760-hour dispatch with 4-phase storage."""
    H = 8760
    total_clean = np.zeros(H)
    for res, pct in resource_pcts.items():
        if pct > 0 and res in supply_profiles:
            total_clean += np.array(supply_profiles[res]) * (pct / 100.0)
    
    # Phase 1-4: Sequential storage dispatch
    for storage_type, pct, eff, dur, window in [
        ('bat4',  battery_dispatch_pct,  BATTERY_EFFICIENCY, 4,    1),
        ('bat8',  battery8_dispatch_pct, BATTERY8_EFFICIENCY, 8,   1),
        ('ldes',  ldes_dispatch_pct,     LDES_EFFICIENCY,    100,  7),
        ('h2',    h2_dispatch_pct,       H2_EFFICIENCY,      1000, 30),
    ]:
        if pct > 0:
            total_clean = _apply_storage(total_clean, demand_norm,
                                          pct, eff, dur, window_days=window)
    
    matched = np.minimum(total_clean, demand_norm * (procurement_pct / 100.0))
    return {\'matched\': matched, \'surplus\': total_clean - matched,
            \'gap\': demand_norm * (procurement_pct / 100.0) - matched}''', font_size=8)
    doc.add_paragraph(
        'Note: Simplified for readability. Full implementation includes per-resource breakdowns, '
        'Numba JIT acceleration (@njit), and SOC carry-across-window logic.',
    )

    doc.add_heading('Code Block 2: Storage Sweep Grid (step1_5_storage_refinement.py)', level=2)
    add_code_block(doc, '''# Gap bucket boundaries (percentage points)
GAP_BUCKET_PP = [5, 10, 25, 50]

# Storage grids (% of annual demand) — union of V1 (near-term) and V2 (2050) caps
FULL_BAT4 = [0, 0.002, 0.005, 0.01, 0.02, 0.03, 0.04, 0.05, 0.06, 0.07, 0.10]
FULL_BAT8 = [0, 0.005, 0.01, 0.02, 0.03, 0.04, 0.06, 0.08, 0.10, 0.15]
FULL_LDES = [0, 0.02, 0.05, 0.1, 0.2, 0.3, 0.5, 0.7, 1.0]
FULL_H2   = [0, 0.3, 1.0]

# Maximum storage for Pass 0 screening
MAX_BAT4 = 0.10   # 0.10% of annual demand
MAX_BAT8 = 0.15   # 0.15% of annual demand
MAX_LDES = 1.0    # 1.0% of annual demand
MAX_H2   = 1.0    # 1.0% of annual demand''', font_size=8)

    doc.add_heading('Code Block 3: Clean Firm Tranching (pipeline_config.py)', level=2)
    add_code_block(doc, '''def compute_clean_firm_tranches(new_cf_twh, iso, firm_lev, ccs_lev, q45,
                                 tx_name=\'Medium\', geo_lev=None,
                                 geo_physics_new_twh=0):
    """Split new clean firm TWh into cost-ordered tranches."""
    # Tranche 1: Nuclear uprates (cheapest, capped per ISO)
    uprate_cap = UPRATE_CAP_TWH.get(iso, 0)
    uprate_twh = min(new_cf_twh, uprate_cap)
    remaining = new_cf_twh - uprate_twh
    
    # Tranche 2: Geothermal (CAISO only, capped at 39 TWh)
    geo_twh = 0
    if iso == \'CAISO\' and geo_lev and remaining > 0:
        geo_cap = GEOTHERMAL_CAP_TWH - geo_physics_new_twh
        geo_twh = min(remaining, max(0, geo_cap))
        remaining -= geo_twh
    
    # Tranche 3: min(nuclear new-build, CCS) - CCS capped per ISO
    ccs_cap = CCS_CAP_TWH.get(iso, 0)
    nuclear_lcoe = NUCLEAR_NEWBUILD_LCOE[firm_lev][iso] + get_tx(\'clean_firm\', tx_name, iso)
    ccs_lcoe = (CCS_LCOE_45Q_ON if q45 == \'1\' else CCS_LCOE_45Q_OFF)[ccs_lev][iso]
    
    if remaining > 0:
        if ccs_lcoe < nuclear_lcoe and ccs_cap > 0:
            ccs_twh = min(remaining, ccs_cap)
            nuclear_twh = remaining - ccs_twh
        else:
            nuclear_twh, ccs_twh = remaining, 0
    else:
        nuclear_twh = ccs_twh = 0
    
    return {\'uprate_twh\': uprate_twh, \'geo_twh\': geo_twh,
            \'nuclear_newbuild_twh\': nuclear_twh, \'ccs_tranche_twh\': ccs_twh}''', font_size=8)

    doc.add_heading('Code Block 4: REC Scarcity Pricing (step6_1_smartargets.py)', level=2)
    add_code_block(doc, '''def compute_rec_price(iso, eligible_pct, year):
    """Scarcity-driven compliance REC price ($/MWh)."""
    acp = ACP_RATES[iso]  # Alternative Compliance Payment cap
    rps_target = get_rps_target_at_year(iso, year)
    vol_adder = VOLUNTARY_DEMAND_ADDER[iso]
    eff_target_pct = (rps_target + vol_adder) * 100.0
    gap = eff_target_pct - eligible_pct  # positive = scarcity
    
    floor = VOLUNTARY_REC_FLOOR[iso]
    k = REC_SCARCITY_K[iso]  # Calibrated per ISO to match 2025 observed
    
    if gap > 0:
        price = acp * (1.0 - np.exp(-k * gap))  # Scarcity ramp toward ACP
    else:
        price = floor + (compliance_2025 - floor) * np.exp(0.20 * gap)  # Decay
    
    return max(floor, min(acp, price))''', font_size=8)

    doc.add_heading('Code Block 5: Wright\'s Law Learning (pipeline_config.py)', level=2)
    add_code_block(doc, '''LEARNING_EXPONENT = 0.6  # Concave ramp

def learning_fraction(year, foak_start, noak_year):
    """Compute Wright\'s Law learning fraction for a given year.
    Returns 0.0 (FOAK) before foak_start, 1.0 (NOAK) after noak_year,
    concave ramp in between: ((year - foak_start) / duration) ** 0.6."""
    if year < foak_start:
        return 0.0
    if year >= noak_year:
        return 1.0
    return ((year - foak_start) / (noak_year - foak_start)) ** LEARNING_EXPONENT

def year_adjusted_cost(foak_cost, noak_cost, year, foak_start, noak_year):
    """Interpolate between FOAK and NOAK costs using Wright\'s Law.
    cost(year) = FOAK * (1 - frac) + NOAK * frac"""
    frac = learning_fraction(year, foak_start, noak_year)
    return foak_cost * (1.0 - frac) + noak_cost * frac''', font_size=8)

    # ======================================================================
    # FOOTER / FINALIZE
    # ======================================================================
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('End of Document')
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Constellation Energy — Commercial Strategy & Analytics\n'
                     'Model Version 1.0.0 | Pipeline v1.0.0 | Base Year 2025')
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    # ======================================================================
    # SAVE
    # ======================================================================
    out_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                            'docs', 'Model_Methodology_Specification.docx')
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f'Saved: {out_path} ({os.path.getsize(out_path) / 1024:.0f} KB)')


if __name__ == '__main__':
    build_document()
